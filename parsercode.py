from __future__ import with_statement
import re
from docx import Document
from docx.enum.text import WD_UNDERLINE
from docx.shared import Pt

# File paths for Ada specifications (.ads) and body (.adb) files.
file_path1 = "filepath1.adb"
file_path3 = "filepath3.adb"
file_path4 = "filepath4.adb"
file_path5 = "filepath5.adb"
file_path6 = "filepath6.adb"
file_path7 = "filepath7.adb"
file_path8 = "filepath8.adb"
file_path9 = "filepath9.adb"
file_path10 = "filepath10.adb"
file_path11 = "filepath11.adb"
file_path12 = "filepath12.adb"
file_path13 = "filepath13.adb"
file_path14 = "filepath14.adb"
file_path15 = "filepath15.adb"

ads_file_path1 = "filepath1.ads"
ads_file_path3 = "filepath3.ads"
ads_file_path4 = "filepath4.ads"
ads_file_path5 = "filepath5.ads"
ads_file_path6 = "filepath6.ads"
ads_file_path7 = "filepath7.ads"
ads_file_path8 = "filepath8.ads"
ads_file_path9 = "filepath9.ads"
ads_file_path10 = "filepath10.ads"
ads_file_path11 = "filepath11.ads"
ads_file_path12 = "filepath12.ads"
ads_file_path13 = "filepath13.ads"
ads_file_path14 = "filepath14.ads"
ads_file_path15 = "filepath15.ads"

common1_file = "file_path_var"
globaldef_file = "file_path_var2"

# Dictionary mapping package names to their corresponding .ads file paths.
Files_With_Procs = {'filepath1':ads_file_path1 ,'filepath3':ads_file_path3 ,'filepath4':ads_file_path4 ,'filepath5':ads_file_path5,'filepath6':ads_file_path6,'filepath7':ads_file_path7,'filepath8':ads_file_path8,'filepath9':ads_file_path9 ,'filepath10':ads_file_path10,'filepath11':ads_file_path11,'filepath12':ads_file_path12,'filepath13':ads_file_path13,'filepath14':ads_file_path14,'filepath15':ads_file_path15}


def find_proc(file_path):
    # This function finds all procedure and function names in a given .ads file.
    # It reads the file line by line and looks for lines starting with 'procedure' or 'function'.
    # The names are extracted and returned in a list.
    procedures = []
    with open(file_path, 'r') as file:
        for line in file:
            s_line = line.strip()  # deleting leading & trailing spaces
            if s_line.lower().startswith('procedure'):
                s_line = re.sub(r'\s*\(\s*', ' ( ', s_line)
                parts = s_line.split()
                p_name = parts[1].split(';')[0]
                procedures.append(p_name)
            elif s_line.lower().startswith('function'):
                s_line = re.sub(r'\s*\(\s*', ' ( ', s_line)
                parts = s_line.split()
                p_name = parts[1]
                procedures.append(p_name)
    return procedures

def files_included(file_path):
    # This function identifies files that are "with-ed" in an Ada package.
    # It parses the 'with...use' statements to find the dependencies of a package.
    included_files = []
    with open(file_path, 'r') as file:
        for line in file:
            s_line = line.strip()
            if s_line.lower().startswith('with') and (';' and 'use' in s_line):
                parts = s_line.split()
                parts[1] = parts[1].split(';')[0]
                included_files.append(parts[1].strip())
    return included_files

def IpsAndOps(file_path, proc_name):
    # This function identifies the global inputs and outputs (variables) used within a procedure's body.
    # It scans the procedure for assignment statements (:=) and function calls to determine which variables are inputs and which are outputs.
    final_ips = {}
    final_out = {}
    inside_proc = False
    with open(file_path, 'r') as file:
        for line in file:
            s_line = line.strip()
            if s_line.startswith('--'):
                continue

            if (s_line.lower().startswith('procedure')) or (s_line.startswith('function')):
                parts = s_line.split()
                current_proc = parts[1].split(';')[0].split('(')[0]
                if current_proc.lower() == proc_name.lower():
                    inside_proc = True
                    inputs, outputs = listOfIpsAndOps(file_path, proc_name)
                    for input in inputs:
                        if isinstance(extract_var_name(clean_variable_name(input)),list):
                            vars = extract_var_name(clean_variable_name(input))
                            for i in range(len(vars)):
                                if '.'.join(vars[:i + 1]) not in final_ips:
                                    dataType = dataTypesOfIps(vars[i], file_path)
                                    final_ips['.'.join(vars[:i+1])] = dataType
                        else:
                            clean_var = clean_variable_name(input)
                            dataType = dataTypesOfIps(clean_var, file_path)
                            final_ips[clean_var] = dataType
                    for output in outputs:
                        if isinstance(extract_var_name(clean_variable_name(output)),list):
                            vars = extract_var_name(clean_variable_name(output))
                            for i in range(len(vars)):
                                if '.'.join(vars[:i+1]) not in final_out:
                                    dataType = dataTypesOfIps('.'.join(vars[:i+1]), file_path)
                                    final_out['.'.join(vars[:i+1])] = dataType
                        else:
                            clean_var = clean_variable_name(output)
                            dataType = dataTypesOfIps(clean_var, file_path)
                            final_out[clean_var] = dataType
            elif inside_proc and s_line.startswith('end'):
                if proc_name.lower() in s_line.lower():
                    break
    return final_ips, final_out


def extract_var_name(var):
    # This utility function extracts variable names from a string.
    # It handles nested record fields (e.g., 'var.field.subfield') and array/function calls (e.g., 'var(index)').
    if '.' in var:
        parts = var.split('.')
        vars = []
        for part in parts:
            variable = part
            if '(' in variable:
                variable = variable.split('(')[0]
            vars.append(variable.strip())
        return vars
    return var.strip()

def clean_variable_name(var):
    # This helper function cleans up a variable name by removing comments, tick marks, and function call parentheses.
    # It's a preprocessing step to isolate the variable name itself.
    if '--' in var:
        var = var.split('--')[0]
    if "'" in var:
        var = var.split("'")[0]
    if '.' not in var and '(' in var:
        var = var.split('(')[0]
    return var.strip()


def listOfIpsAndOps(file_path, proc):
    # This function extracts all variables that are read from (inputs) or written to (outputs)
    inputs = []
    outputs = []
    exc_key=['if','not','return','for','loop','or']
    with open(file_path, 'r') as file:
        isFound = False
        inBody = False
        insideifBlock = False
        ifblock = []
        checkArray = ['true', 'false']
        symbols = ['/=', '>=','<=', '+', '=', '/', '<', '>', '-', '*']
        loopVar = 'null'
        contSym = False
        for line in file:
            containSymbol = False
            if not isFound:
                s_line = line.strip()
                mSym = " ( "
                s_line = s_line.replace("(", mSym)
                if s_line.lower().startswith('procedure'):
                    parts = s_line.split()
                    if parts[1] == proc and not parts.__contains__(";") :
                        isFound = True
                        continue
                if s_line.lower().startswith('function'):
                    parts = s_line.split()
                    if parts[1] == proc:
                        isFound = True
                        continue
            else:
                line = line.split('--')[0].strip().replace('(', ' ( ').replace(')', ' ) ').replace('=>',' => ')
                if line.startswith("--"):
                    continue
                if line.lower().startswith('begin'):
                    inBody = True
                if line.lower().startswith('if'):
                    insideifBlock = True
                if line.lower().startswith('elsif'):
                    insideifBlock = True
                if insideifBlock and line.lower().__contains__('then'):
                    ifblock.append(line)
                    insideifBlock = False
                if inBody and line.startswith('end') and re.search(r'\b'+re.escape(proc)+r'\b',line,re.IGNORECASE):#line.lower().__contains__(proc):
                    break
                if inBody and line.startswith('for'):
                    loopVar = line.split(' ')[1]
                
                if inBody and insideifBlock:
                    ifblock.append(line)
                if 'not ' in line:
                     modSym = " ) "
                     modSy=" ( "
                     line = line.replace(')', modSym).replace("(",modSy)
                     parts = line.split()
                     if 'not' in parts:
                         idxNot = line.index('not ')
                         if idxNot+1<len(parts):
                             value=parts[idxNot+1]
                     if (
                         value.lower() not in exc_key and
                         value.lower() not in inputs and
                         value.lower() not in checkArray):
                             inputs.append(value)

                for sym in symbols:
                    if line.__contains__(sym):
                        if sym == '=':
                            if line[line.index(sym)-1] == ":" :
                                continue
                        containSymbol = True
                        break

                if contSym:
                    tem_part = line.split(';')[0]
                    if not tem_part.__contains__(" ") and inputs.count(tem_part) == 0:
                        inputs.append(tem_part)
                    contSym = False

                if inBody and not insideifBlock and ':=' in line:
                    left, right = map(str.strip, line.split(':=',1))
                    right = right.split(';')[0]

                    if outputs.count(left) == 0: # Not considering repeated outputs
                        outputs.append(left)

		    if right.__contains__('('):
                        teSym = " ( "
                        tempRht = right.replace('(', teSym)
                        tempRht = tempRht.split()
                        if tempRht[0] == "(":
                            temp1 = tempRht[1]
                            if listOfTheProcs(temp1):
                                continue
                            if len(tempRht) > 3:
                                temp2 = tempRht[3]
                                if listOfTheProcs(temp2):
                                    continue
                        else:
                            temp1 = tempRht[0]
                            if listOfTheProcs(temp1):
                                continue
                            if len(tempRht) > 2:
                                temp2 = tempRht[2]
                                if listOfTheProcs(temp2):
                                    continue

                    if listOfTheProcs(right):
                        continue


                    for sym in symbols:
                       if right.__contains__(sym):
                           if sym == '=':
                               if line[line.index(sym) - 1] == ":":
                                   break
                           containSymbol = True
                           break

                    if right == loopVar: # For omitting the constant values
                        loopVar = 'null'
                        continue
                    if right == "":
                        continue
                    if '0' <= right[0] <= '9' or right[0] == "'":
                        continue
                    if re.search(r'[a-zA-Z]', right) and containSymbol == False and checkArray.count(right) == 0: # For not considering the keywords
                        if inputs.count(right) == 0  and outputs.count(right) == 0: # Considering only the inputs which are not outputs
                            inputs.append(right)
                    tem_parts = line.split()
                    val = tem_parts[len(tem_parts)-1]
                    if symbols.__contains__(val): #For checking the incomplete eqns
                        contSym = True

                if containSymbol:
                    tokens = split_by_symbols(line, symbols)
                    for i in range(1, len(tokens) - 1):
                        if tokens[i] in symbols:
                            bef = tokens[i - 1]
                            aft = tokens[i + 1]
                            if outputs.count(bef) == 0 and inputs.count(bef) == 0 and re.search(r'[a-zA-Z]', bef):
                                inputs.append(bef)
                            if outputs.count(aft) == 0 and inputs.count(aft) == 0 and 'a' <= aft[0].lower() <= 'z' and re.search(r'[a-zA-Z]', aft) and checkArray.count(aft) == 0:
                                inputs.append(aft)

    for line in ifblock:
        line = line.replace(')', ' ').replace('(', ' ')
        parts = line.split()
        for part in parts:
            if is_key(part):
                continue
            if is_dataType(part):
                continue
            if part in symbols:
                continue
            if re.search(r'^[A-Za-z][A-Za-z0-9_]*$', part):
                inputs.append(part)

        return inputs, outputs

def split_by_symbols(line, symbols):
    # This function is designed to split a line of code by various symbols.
    # It adds spaces around symbols to ensure they are treated as separate tokens during splitting.
    line = line.replace(';', ' ')
    for sym in symbols:
        if line.__contains__(sym):
            if sym == '=':
                if line[line.index(sym)-1] == ":" or line[line.index(sym)-1] == ">" or line[line.index(sym)-1] == "/":
                    continue
            if sym == '-':
                if line.index(sym) + 1 != len(line):
                    if line[line.index(sym) + 1] == '-':
                        break
            modSym = ' ' + sym + ' '
            line = line.replace(sym, modSym)
    parts = line.split()

    mod_line = ""
    for part in parts:
        if part.startswith('('):
            part = part[1:len(part)]
        if part.endswith(')'):
            part = part[0:len(part) - 1]
        mod_line =  mod_line + " " + part
    return mod_line.split()

def dataTypesOfIps(input, file_path):
    common1_file_path = "file_path_var1"
    with open(common1_file_path, 'r') as file:
        for line in file:
            line = line.strip()
            if line.lower().__contains__(input.lower()) and line.__contains__(':'):
                right = (line.split(':')[1]).split(';')[0].strip()
                return right

    globaldef_file_path = "file_path_var2"
    with open(globaldef_file_path, 'r') as file2:
        for line in file2:
            line = line.strip() # for getting rid of leading and trailing spaces
            if line.lower().__contains__(input.lower()) and line.__contains__(':'): # For getting rid of case-sensitiviy in python
                right = (line.split(':')[1]).split(';')[0].strip()
                return right

    with open(file_path, 'r') as file3:
        for line in file3:
            line = line.strip()
            if line.lower().__contains__(input.lower()) and line.__contains__(':') and not line[line.index(':')-1] == "=": # For getting rid of case-sensitiviy in python
                right = (line.split(':')[1]).split(';')[0].strip()
                return right
    return 'null'


def listOfTheProcs(proc):
    # This function checks if a given name corresponds to a procedure or function in any of the specified .ads files.
    # It is used to identify calls to other modules within the current procedure's body.
    
    files = [ads_file_path1, ads_file_path3, ads_file_path4, ads_file_path5, ads_file_path6, ads_file_path7, ads_file_path8, ads_file_path9, ads_file_path10, ads_file_path11, ads_file_path12, ads_file_path13, ads_file_path14, ads_file_path15]#, file_path16]
    for file in files:
        result = checkProc(file,proc)
        if result == True:
            return True
    return False

def checkProc(file_name, proc):
    # A helper function that searches a single file for a procedure or function by name.
    # It returns True if the procedure exists and False otherwise.
    procs = []
    with open(file_name, 'r') as file:
        for line in file:
            s_line = line.strip()
            if s_line.lower().startswith('procedure'):
                mod_line = s_line.replace('(', ' ')
                parts = mod_line.split()
                p_name = parts[1].split(';')[0]
                procs.append(p_name)
            elif s_line.lower().startswith('function'):
                mod_line = s_line.replace('(', ' ')
                parts = mod_line.split()
                p_name = parts[1]
                procs.append(p_name)
    return procs.__contains__(proc)

def getName(file_path):
    # Extracts the package name from a given file path.
    # It assumes the file name is the package name.
    parts = file_path.split('/')
    name = parts[len(parts)-1]
    name = name.split('.')[0]
    print("Package: " + name)
    print("")
    return name

def local_variables(file_path,proc_name):
    # This function finds all locally declared variables within a specific procedure or function.
    # It looks for variable declarations between the 'is' and 'begin' keywords.
    inside_proc = False
    inside_local_block = False
    local_var_name = {}
    with open(file_path, 'r') as file:
        for line in file:
            if ('(' or ')') in line:
                line = line.replace('(', ' ( ')
                line = line.replace(')', ' ) ')
            line = line.strip()
            s_line = line.split('--')[0].strip()
            if not s_line:
                continue
            if not inside_proc and re.match(r'\b(procedure|function)\s+' + r'\b'+re.escape(proc_name)+r'\b', s_line, re.IGNORECASE):
                inside_proc = True
                if inside_proc and re.search(r'\bis\b',s_line,re.IGNORECASE):
                    inside_local_block = True
                    continue
            if inside_proc and not inside_local_block and re.search(r'\bis\b',s_line,re.IGNORECASE):
                inside_local_block=True
                continue
            if inside_proc and inside_local_block and re.match(r'\bbegin\s*',s_line,re.IGNORECASE):
                break
            if inside_proc and inside_local_block and ':' in s_line and not ':='  in s_line :
                s_line = line.split('--')[0].strip()
                if s_line:
                    parts=s_line.split(':')
                    varname = parts[0].strip()
                    datatype = parts[1].strip()
                    datatype = datatype.split(';')[0]
                    local_var_name[varname] = datatype
    return local_var_name

def is_argument(file_path,proc_name):
    # This function parses a procedure or function declaration to identify its arguments and their modes (in, out, in out).
    # It also captures the return type for functions.
    args_in={}
    args_out={}
    args_inout={}
    return_type=None
    inside_proc = False
    param_line=[]
    with open(file_path, 'r') as file:
        for line in file:
            s_line = line.strip()
            if s_line.startswith('--'):
                continue
            if not inside_proc and re.match(r'\b(procedure|function)\s+' + r'\b'+re.escape(proc_name)+r'\b', s_line, re.IGNORECASE):#Regex  to prevent spillover
                inside_proc = True
                param_line.append(s_line)
            elif inside_proc:
                param_line.append(s_line)
            if inside_proc and s_line.startswith('begin'):
                break
    param=' '.join(param_line)
    if 'function' in param.lower():
        match=re.search(r'return\s*([A-Za-z0-9_]*)',param,re.IGNORECASE)
        if match:
            return_type=match.group(1).strip()
    if '(' in param and ')' in param:
        start=param.index('(')+1
        end=param.rindex(')')
        param_block=param[start:end].strip()
    else:
        return args_in,args_out,args_inout,return_type
    param_list=[p.strip() for p in param_block.split(';')if p.strip()]
    for params in param_list:
        if ':' not in params:
            continue
        left,right=params.split(':',1)
        names=[left.strip()]
        right_parts=right.lower().split()
        mode='in'#default
        if 'in' in right_parts and 'out' in right_parts:
            mode='inout'
        elif 'out' in right_parts:
            mode='out'
        elif 'in' in right_parts:
            mode='in'
        dataType=next((word for word in right_parts if word not in['in','out']),'unknown').upper()
        for name in names:
            if mode=='in':
                args_in[name] = dataType
            elif mode=='out':
                args_out[name] = dataType
            elif mode=='inout':
                args_inout[name] = dataType
    return args_in,args_out,args_inout,return_type

def is_constant(file_path):
    # Identifies and extracts constants declared in a file.
    # It looks for lines containing 'constant', ':=', ':', and ';', then parses the name, type, and value.
    constant=[]
    with open(file_path, 'r') as file:
        for line in file:
            line = line.strip()
            if line.startswith('--'):
                continue
            if 'constant' in line.lower() and ':=' in line and ':'in line and ';' in line:#check for 'constant',':=',':' and ';'
                try:
                    line=line.split('--')[0].strip()#remove the leading and trailing spaces along with comment
                    name_part,rest=line.split(':',1)
                    name_part=name_part.strip()#removes whitespaces
                    rest=rest.split('constant',1)
                    type_part=rest[1].split(':=',1)[0].strip() #to get the type which is before ':='
                    value_part=rest[1].split(':=',1)[1].strip().rstrip(';')#to get the value after ':='
                    constant.append((name_part,type_part,value_part))
                except Exception as e:
                    print(f"Error: {line}\n{e}")#shows if any error and the type of error
    return constant

def ext_proc(file_path,proc_name):
    # Extracts the complete body of a procedure, from the 'begin' keyword to the 'end <procedure_name>'.
    # This is useful for later analysis, like finding constants used within the body.
    inside_proc = False
    inside_body = False
    body_line = []
    with open(file_path, 'r') as file:
        for line in file:
            s_line = line.split('--')[0].strip()
            if not s_line:
                continue
            if not inside_proc and re.match(r'\b(procedure|function)\s+' + r'\b'+re.escape(proc_name)+r'\b', s_line, re.IGNORECASE):  # Regex  to prevent spillover
                inside_proc = True
                continue
            if inside_proc and not inside_body and re.search(r'\bbegin\b', s_line, re.IGNORECASE):
                inside_body = True
            if inside_body:
                body_line.append(line)
                if s_line.lower().startswith('end') and proc_name.lower() in s_line.lower():
                    break
    return '\n'.join(body_line)

def match_const(proc_text,constant):#which match the constant name by checking cname and proc_text
    # This function finds all constants from a given list that are used within a procedure's body.
    # It checks for a case-insensitive match of the constant name within the procedure text.
    used_constant=[]
    proc_text=proc_text.lower()
    for cname,ctype,cvalue in constant:
        pattern=r'\b'+re.escape(cname.lower())+r'\b'#convert cname to lower case to easily match with proc_text
        if re.search(pattern,proc_text):
            used_constant.append((cname,ctype,cvalue))
    return used_constant

def is_in_Const(name,const_list):
    # A simple helper function to check if a variable name is present in a list of constants.
    name=name.lower()
    return any(cname.lower()==name for cname,ctype,cvalue in const_list)

def fetch_modules(file_path, proc_name, To_list_of_procs):
    # This function identifies which modules (procedures or functions) are called from within a specific procedure's body.
    # It compares words in the procedure body against a provided list of known procedures to find matches.
    inside_proc = False
    inside_body=False
    List_of_procs = []
    with open(file_path, 'r') as file:
        for line in file:
            s_line = line.split('--')[0].replace('(',' ( ').replace(')',' ) ').strip()
            s_line=s_line.split(';')[0].strip()
            if not s_line:
                continue
            if not inside_proc and  re.match(r'\b(procedure|function)\s+'+r'\b'+re.escape(proc_name)+r'\b',s_line,re.IGNORECASE):
                inside_proc = True
                continue
            if inside_proc and re.match(r'\bbegin\b',s_line,re.IGNORECASE):
                inside_body = True
                continue
            if inside_proc and re.match(r'\bend\s+'+r'\b'+re.escape(proc_name)+r'\b',s_line,re.IGNORECASE):#For getting rid of case-sensitivity in python
                break
            if not inside_body:
                continue
            if inside_body:
                parts = s_line.split()
                for i in parts:
                    for j in To_list_of_procs:
                        if i.lower() == j.lower():
                            List_of_procs.append(i)
        List_of_procs = set(List_of_procs)
        list(List_of_procs)
    return List_of_procs

def is_key(var_name):
    # Checks if a given string is an Ada programming language keyword.
    keywords =  { 'for', 'return', 'not', 'while', 'loop', 'then', 'function', 'end','false','true','elsif', 'exit', 'and','if','or','case','when','pragma','raise','delay','sqrt','abs'}
    if var_name.lower() in keywords:
        return True
    return False

def is_dataType(var_name):
    # Checks if a given string is one of the simple, primitive data types.
    dataTypes = {'INTEGER', 'FLOAT', 'BOOLEAN', 'CHARACTER', 'STRING'}
    if var_name in dataTypes:
        return True
    return False

if __name__ == '__main__':
    doc = Document()
    para = doc.add_paragraph()
    file_path =file_path1
    procedures = find_proc(file_path)
    common_const=is_constant(common1_file)
    global_const=is_constant(globaldef_file)
    all_const=common_const+global_const #combine the files to easily check each procedure
    file_name = getName(file_path)#to get the package name
    List_of_procs = []
    if file_name in Files_With_Procs:
        List_of_procs = find_proc(Files_With_Procs[file_name])
    run=para.add_run('Package: '+file_name+'\n\n')
    run.font.size = Pt(20)
    run.bold = True
    run.underline =  WD_UNDERLINE.DOUBLE
    run = para.add_run('\n<designer to enter package overview>\n')
    run.font.size = Pt(15)
    run.bold = True
    run.centered = True
    included_files = files_included(file_path)
    for included_file in included_files:
        if included_file in Files_With_Procs:
            List_of_procs += find_proc(Files_With_Procs[included_file])
    para.add_run('\n\n')
    run=para.add_run('Modules in the Package : \n')
    run.font.size = Pt(15)
    run.bold = True
    run.underline = True
    for proc_name in procedures:
        para.add_run(proc_name+'\n')
    for proc_name in procedures:
        print(f"\nPROCEDURE : {proc_name}")
        global_inputs = {}
        local_var = local_variables(file_path,proc_name)
        global_outputs = {}
        ips, ops = IpsAndOps(file_path, proc_name)
        args_in,args_out, args_inout,return_type=is_argument(file_path,proc_name)
        for var in ips:
            clean_var=clean_variable_name(var)
            if ('(' or ')') in clean_var:
                clean_var = clean_var.replace('(', ' ( ')
                clean_var = clean_var.replace(')', ' ) ')
            if is_key(clean_var):
                continue
            if clean_var.startswith('--'):
                continue
            if is_in_Const(clean_var, all_const):
                continue
            if is_dataType(clean_var):
                continue
            if clean_var.split('.', 1)[0] in local_var:
                local_var[var] = local_dict[var]
            localflag = False
            for i in local_var:
                if clean_var.lower() == i.lower():
                    localflag = True
            if not localflag:
                global_inputs[clean_var] = dtype
        for var in ops:
            clean_var = clean_variable_name(var)
            if ('(' or ')') in clean_var:
                clean_var = clean_var.replace('(', ' ( ')
                clean_var = clean_var.replace(')', ' ) ')
            if is_key(clean_var):
                continue
            if clean_var.startswith('--'):
                continue
            if is_in_Const(clean_var, all_const):
                continue
            if is_dataType(clean_var):
                continue
            if clean_var.split('.', 1)[0] in local_var:
                local_var[var] = local_dict[var]
            localflag = False
            for i in local_var:
                if clean_var.lower() == i.lower():
                    localflag = True
            if not localflag:
                global_outputs[clean_var] = dtype
        proc_text=ext_proc(file_path,proc_name)
        consts=match_const(proc_text,all_const)

        modules = fetch_modules(file_path, proc_name, List_of_procs)
        run=para.add_run(f"\n\nPROCEDURE : {proc_name}\n")
        run.font.size = Pt(16)
        run.bold = True
        run.underline = True
        run = para.add_run("\nINPUTS: \n")
        run.font.size = Pt(14)
        run.bold = True
        run.underline = True
        run = para.add_run("\nARGUMENTS:\n")
        run.bold = True
        if args_in:
            for arg in args_in:
                para.add_run(arg + ': ' + args_in[arg]+'\n')
        else:
            para.add_run('None')
        run = para.add_run(" \nGLOBAL VARIABLES:")
        run.bold = True
        if global_inputs:
            for var, dtype in global_inputs.items():
                if var.lower() not in ['true', 'false']:
                    if var in args_in:
                        continue
                    if var in args_out:
                        continue
                    if var in args_inout:
                        continue
                    clean_type = re.split(r'\s+(range)\s+', dtype, flags=re.IGNORECASE)[0]
                    if '.' in var:
                        para.add_run('\n'+('\t'*var.count('.'))+f"{var} : {clean_type}")
                    else:
                        para.add_run(f"\n{var} : {clean_type}")
        else:
            para.add_run("\nNone")
        run=para.add_run("\n\nOUTPUTS: \n")
        run.font.size = Pt(14)
        run.bold = True
        run.underline = True
        run = para.add_run("\nARGUMENT:\n")
        run.bold = True
        if args_out:
            for arg in args_out:
                para.add_run(arg + ': ' + args_out[arg]+'\n')
        else:
            para.add_run('None')
        run = para.add_run(" \nGLOBAL VARIABLES:")
        run.bold = True
        if global_outputs:
            for var, dtype in global_outputs.items():
                if var in args_in:
                    continue
                if var in args_out:
                    continue
                if var in args_inout:
                    continue
                clean_type = re.split(r'\s+(range)\s+', dtype, flags=re.IGNORECASE)[0]
                if '.' in var:
                    para.add_run('\n'+('\t'*var.count('.'))+f"{var} : {clean_type}")
                else:
                    para.add_run(f"\n{var} : {clean_type}")
        else:
            para.add_run("\nNone")
        run = para.add_run("\n\nARGUMENT InOuts:")
        run.font.size = Pt(13)
        run.bold = True
        if args_inout:
            for arg in args_inout:
                para.add_run(arg + ': ' + args_inout[arg]+'\n')
        else:
            para.add_run('None')
        run = para.add_run("\nRETURN TYPE:")
        para.add_run(f"\n{return_type}\n")
        run.font.size = Pt(13)
        run.bold = True
        run.underline = True
        run = para.add_run("\nLOCAL VARIABLES: \n")
        run.font.size = Pt(13)
        run.bold = True
        run.underline = True
        if local_var:
            for var, dtype in local_var.items():
                if var.lower() not in ['true', 'false']:
                    clean_type = re.split(r'\s+(range)\s+', dtype, flags=re.IGNORECASE)[0]
                    clean_type = str(clean_type)
                    para.add_run(('\t'*var.count('.'))+f" {var} : {clean_type}\n")
            para.add_run('')
        else:
            para.add_run("None\n")
        run = para.add_run("\nMODULES CALLED:\n")
        run.font.size = Pt(13)
        run.bold = True
        run.underline = True
        if modules:
            for module in modules:
                para.add_run(module+'\n')
        else:
            para.add_run('None\n')
        run = para.add_run("\nCONSTANT: ")
        run.font.size = Pt(13)
        run.bold = True
        run.underline = True
        para = doc.add_paragraph()
        header = ['Name', 'Type', 'Value']
        table = doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        for i,header in enumerate(header):
            hdr_cells[i].text = header
        if consts:
            for cname, ctype, cvalue in consts:
                localflag = False
                for i in local_var:
                    if cname.lower() == i.lower():
                        localflag = True
                        break
                if not localflag:
                    const = [cname, ctype, cvalue]
                    row_cells = table.add_row().cells
                    for i, const in enumerate(const):
                        row_cells[i].text = const
        else:
            const = ['None', '-', '-']
            row_cells = table.add_row().cells
            for i, const in enumerate(const):
                row_cells[i].text = const
        table.style = 'Table Grid'

        para = doc.add_paragraph()
        run = para.add_run("\nDESCRIPTION\n")
        run.font.size = Pt(15)
        run.bold = True
        run.underline = True
        run = para.add_run("\n\n<designer to enter procedure overview>\n\n\n")
        run.bold = True
        run.centered = True
        run = para.add_run("\nBODY OF PROCEDURE:\n")
        run.font.size = Pt(13)
        run.bold = True
        run.underline = True
        para.add_run(proc_text)
        print(common_global_dict)
    doc.save('Parser.docx')
